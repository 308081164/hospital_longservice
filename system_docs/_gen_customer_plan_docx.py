"""Generate Word/PDF from customer development plan markdown."""
from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).parent
MD_FILE = BASE / "冠唐云仓库-货品批量导入-客户开发计划.md"
DOCX_FILE = BASE / "冠唐云仓库-货品批量导入-客户开发计划.docx"
PDF_FILE = BASE / "冠唐云仓库-货品批量导入-客户开发计划.pdf"


def set_cn_font(run, name="微软雅黑", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, style=None, bold=False, size=11, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_cn_font(run, size=size)
    run.bold = bold
    return p


def parse_table_lines(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cn_font(run, size=10)
                    if i == 0:
                        run.bold = True


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("# "):
            add_paragraph(doc, stripped[2:], bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        if stripped.startswith("## "):
            add_paragraph(doc, stripped[3:], bold=True, size=14)
            i += 1
            continue
        if stripped.startswith("### "):
            add_paragraph(doc, stripped[4:], bold=True, size=12)
            i += 1
            continue
        if stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], size=10)
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table_lines(table_lines))
            continue
        if re.match(r"^\d+\.\s", stripped):
            add_paragraph(doc, stripped, size=11)
            i += 1
            continue
        if stripped.startswith("- "):
            add_paragraph(doc, "• " + stripped[2:], size=11)
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*"):
            add_paragraph(doc, stripped.strip("*"), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        add_paragraph(doc, stripped, size=11)
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def docx_to_pdf(docx_path: Path, pdf_path: Path):
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"Created: {pdf_path}")
        return True
    except Exception as e:
        print(f"PDF conversion failed: {e}", file=sys.stderr)
        return False


if __name__ == "__main__":
    md_to_docx(MD_FILE, DOCX_FILE)
    docx_to_pdf(DOCX_FILE, PDF_FILE)
