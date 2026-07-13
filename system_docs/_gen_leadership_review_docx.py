# -*- coding: utf-8 -*-
"""Generate Word/PDF from leadership review delivery txt."""
from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).parent
TXT_FILE = BASE / "领导审查交付说明.txt"
DOCX_FILE = BASE / "领导审查交付说明.docx"
PDF_FILE = BASE / "领导审查交付说明.pdf"

SEP_RE = re.compile(r"^=+$")
TABLE_RULE_RE = re.compile(r"^-{4,}")


def set_cn_font(run, name="微软雅黑", size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_run_paragraph(doc, text, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold)
    return p


def is_section_title(line, lines, idx):
    if not SEP_RE.match(line.strip()):
        return False
    if idx + 1 >= len(lines):
        return False
    nxt = lines[idx + 1].strip()
    if not nxt or SEP_RE.match(nxt):
        return False
    if idx + 2 >= len(lines):
        return False
    return SEP_RE.match(lines[idx + 2].strip()) is not None


def is_subsection(line):
    s = line.strip()
    if re.match(r"^（[一二三四五六七八九十]+）", s):
        return True
    if re.match(r"^第[一二三四五六七八九十]+，", s):
        return True
    return False


def split_table_row(line):
    parts = re.split(r"\s{2,}", line.strip())
    return [p.strip() for p in parts if p.strip()]


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    norm = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(norm), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(norm):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cn_font(run, size=10, bold=(i == 0))
    doc.add_paragraph()


def txt_to_docx(txt_path: Path, docx_path: Path):
    lines = txt_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    i = 0
    title_lines = []
    while i < len(lines) and not lines[i].strip():
        i += 1
    while i < len(lines) and lines[i].strip() and not SEP_RE.match(lines[i].strip()):
        title_lines.append(lines[i].strip())
        i += 1
    if title_lines:
        add_run_paragraph(
            doc,
            title_lines[0],
            size=20,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=4,
        )
    if len(title_lines) > 1:
        add_run_paragraph(
            doc,
            title_lines[1],
            size=16,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=8,
        )
    for meta in title_lines[2:]:
        add_run_paragraph(
            doc,
            meta,
            size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=4,
        )
    if title_lines:
        doc.add_paragraph()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if is_section_title(stripped, lines, i):
            heading = lines[i + 1].strip()
            add_run_paragraph(doc, heading, size=14, bold=True, space_after=8)
            i += 3
            continue

        if SEP_RE.match(stripped):
            i += 1
            continue

        if TABLE_RULE_RE.match(stripped):
            i += 1
            continue

        if stripped.startswith("序号") or stripped.startswith("项目名称") or ("计划交付时间" in stripped and "当前状态" in stripped):
            header = split_table_row(stripped)
            i += 1
            if i < len(lines) and TABLE_RULE_RE.match(lines[i].strip()):
                i += 1
            body_rows = [header]
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line:
                    break
                if SEP_RE.match(row_line) or is_section_title(row_line, lines, i):
                    break
                if TABLE_RULE_RE.match(row_line):
                    i += 1
                    continue
                if row_line.startswith("医院计费") or row_line.startswith("以上材料"):
                    break
                cells = split_table_row(row_line)
                if cells and re.match(r"^\d+$", cells[0]):
                    body_rows.append(cells)
                    i += 1
                    continue
                if len(cells) >= 2 and "2026" in row_line:
                    body_rows.append(cells)
                    i += 1
                    continue
                break
            add_table(doc, body_rows)
            continue

        if is_subsection(stripped):
            add_run_paragraph(doc, stripped, size=12, bold=True, space_after=6)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_run_paragraph(doc, stripped, size=11, space_after=6)
            i += 1
            continue

        if stripped.startswith("文档版本") or stripped.startswith("编制单位") or stripped.startswith("联系方式"):
            add_run_paragraph(
                doc,
                stripped,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=2,
            )
            i += 1
            continue

        add_run_paragraph(doc, stripped, size=11, space_after=6)
        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def docx_to_pdf(docx_path: Path, pdf_path: Path):
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        print(f"Created: {pdf_path}")
        return True
    except Exception as e:
        print(f"PDF conversion failed: {e}", file=sys.stderr)
        return False


if __name__ == "__main__":
    txt_to_docx(TXT_FILE, DOCX_FILE)
    docx_to_pdf(DOCX_FILE, PDF_FILE)

