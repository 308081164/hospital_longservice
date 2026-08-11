#!/usr/bin/env python3
"""一期优先医院：特色导出功能规则清单（Word，供客户核对）。"""

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

MANIFEST_PATH = ROOT / "backend/src/main/resources/billing-seeds/billing-rules-manifest.json"
EXPORT_SEED = ROOT / "backend/src/main/resources/billing-seeds/phase-export-rules-20260723.json"
DEPT_SEED = ROOT / "backend/src/main/resources/billing-seeds/phase-export-dept-split-20260728.json"
OUT_DOCX = ROOT / "测试用例/一期优先医院特色导出规则清单.docx"
OUT_MD = ROOT / "测试用例/一期优先医院特色导出规则清单.md"

# 来源：测试用例/优先医院对齐TODO.md「逐院导出文件类型」表（2026-07-29）
EXPORT_FILE_TYPES: dict[str, list[str]] = {
    "ZYY-D1": ["账单", "结款函", "分科室汇总表", "物流分摊表"],
    "ZY3-DIANLI": ["账单", "结款函", "器械把数表"],
    "GUOYAO-MAIN": ["账单", "结款函"],
    "GUOYAO-2": ["账单", "结款函"],
    "GUOYAO-3": ["账单", "结款函"],
    "HRB-2ND": ["账单", "结款函"],
    "HRB-WY": ["账单", "结款函", "分科室汇总表", "价格汇总表", "器械把数表", "总汇总表"],
    "HRB-WY-EM": ["账单", "结款函", "总汇总表"],
    "XINFA-HSZ": ["账单", "结款函"],
    "SHENG-YY-NG": ["账单", "结款函", "价格汇总表", "器械把数表", "物流分摊表"],
    "SHENG-YY-XF": ["账单", "结款函", "价格汇总表", "器械把数表", "物流分摊表"],
    "ZUYAN-NG": ["账单", "结款函", "价格汇总表"],
    "ZUYAN-SF": ["账单", "结款函", "价格汇总表"],
    "ZUYAN-XA": ["账单", "结款函", "价格汇总表"],
    "NG-FUCHAN": ["账单", "结款函"],
    "SHKF-YY": ["账单", "结款函"],
    "DAOWAI-RM": ["账单", "结款函"],
    "TAIPING-RM": ["账单", "结款函"],
    "SANJING-SB": ["账单", "结款函"],
    "VICTORIA": ["账单", "结款函"],
    "JIUZHOU-FK": ["账单", "结款函"],
    "HULAN-HSZ": ["账单", "结款函"],
    "HULAN-TCM": ["账单", "结款函"],
    "ZYY-D2-NG": ["账单", "结款函", "价格汇总表", "器械把数表"],
    "ZYY-D2-HN": ["账单", "结款函", "价格汇总表", "器械把数表"],
    "RENSHENG": ["账单", "结款函"],
    "HRB-HX-EYE": ["账单", "结款函"],
    "BINGCHENG-YM": ["账单", "结款函"],
    "XF-ZYY": ["账单", "结款函"],
    "WJ-HLJ-ZD": ["账单", "结款函"],
    "YUEMEI-FH": ["账单", "结款函"],
    "ERYY-NG": ["账单", "结款函"],
    "ERYY-SB": ["账单", "结款函"],
    "HULAN-RM": ["账单", "结款函"],
    "HRB-HSZ": ["账单", "结款函"],
    "HRB-HIT": ["账单", "结款函"],
    "HRB-HEU": ["账单", "结款函"],
}

BILL_STRATEGY_PLAIN = {
    "standard_bill": "标准账单格式（与多数医院相同）",
    "guoyao_bill": "国药专用：按「汽轮机」算法折算包数后出账",
    "daowai_bill": "道外区人民医院专用账单格式",
    "sheng_er_bill": "省二医院（南岗）专用账单格式",
}

BILL_LAYOUT_PLAIN = {
    "dept_split": "按科室分成多个工作表（每个科室一页）",
    "combined": "全部明细合在一个工作表里",
    "auto": "自动：一般一个表；明细特别多时可能按科室分页",
    "fuyi_extended_11col": "附一专用：11 列宽表头",
}


def set_doc_font(doc: Document, name: str = "微软雅黑") -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def merge_templates(export_seed: dict, dept_seed: dict) -> dict[str, list[dict]]:
    by_code: dict[str, list[dict]] = {}
    for src in (export_seed, dept_seed):
        for o in src.get("exportTemplateOverrides", []):
            by_code.setdefault(o["code"], []).extend(o.get("templates", []))
    return by_code


def describe_template(t: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    ttype = t.get("type", "")
    if ttype == "bill":
        lines.append("账单表")
    elif ttype == "settlement":
        lines.append("结款函")
    else:
        lines.append(str(ttype))
    strat = t.get("strategyKey", "")
    if strat:
        lines.append(f"  格式：{BILL_STRATEGY_PLAIN.get(strat, strat)}")
    cm = t.get("columnMapping") or {}
    layout = cm.get("billLayout")
    if layout:
        lines.append(f"  排版：{BILL_LAYOUT_PLAIN.get(layout, layout)}")
    if cm.get("billColumnLayout"):
        lines.append(f"  列版式：{BILL_LAYOUT_PLAIN.get(cm['billColumnLayout'], cm['billColumnLayout'])}")
    if cm.get("d8DisplaySource") == "hospitalName":
        lines.append("  明细里 D 列显示医院名称（不是规则名）")
    if cm.get("settlementDiscountRows"):
        lines.append("  结款函里单独列出折扣行")
    remove = cm.get("removeColumns")
    if remove:
        lines.append(f"  删掉这些列：{'、'.join(remove)}")
    keep = cm.get("keepColumns")
    if keep:
        cols = "、".join(keep[:10])
        if len(keep) > 10:
            cols += "…"
        lines.append(f"  只保留这些列：{cols}")
    return lines


def plain_rule_line(rule: dict[str, Any]) -> str:
    """客户可读的单条规则说明。"""
    name = rule.get("name") or ""
    kws = rule.get("keywords") or []
    kw_text = "、".join(str(k) for k in kws[:6])
    if len(kws) > 6:
        kw_text += "等"
    price = rule.get("price")
    rt = rule.get("ruleType", "")
    cj = rule.get("conditionsJson") or ""
    export_only = "exportApply" in cj

    if rt == "FIXED_PRICE" and price is not None:
        body = f"包名里带「{kw_text}」的，按 **{price:g} 元一包**收"
    elif rt == "PRICE_PER_INSTRUMENT" and price is not None:
        body = f"包名里带「{kw_text}」的，按 **{price:g} 元一件**收"
    elif rt == "FOLD":
        th, fr = rule.get("threshold"), rule.get("foldRatio")
        body = f"包名里带「{kw_text}」的，**{th} 件算 {fr} 件**来收钱" if th and fr else f"包名里带「{kw_text}」"
    else:
        body = f"包名里带「{kw_text}」— {name}"

    if export_only:
        body += "（**只在导出账单 Excel 时**用这个价，不影响系统对账）"
    return f"{name}：{body}"


def plain_discount(text: str | None) -> str | None:
    if not text:
        return None
    t = text
    t = t.replace("export_only", "仅导出账单时")
    t = t.replace("phase5-batch-c", "")
    t = t.replace("（", "（").strip()
    if "2+把75%" in t or "75%阶梯" in t:
        return "同一把器械满 2 件及以上，导出账单时再打 75 折"
    if "结款7折" in t:
        return "结款函里灭菌费按 7 折结算"
    if "账单7折" in t:
        return "导出账单时整体打 7 折"
    if "结款灭菌9折" in t:
        return "结款函里灭菌费按 9 折结算"
    if "加急" in t:
        return "结款函里单独列加急费、设备抵扣（见新发红十字等院配置）"
    return t.strip(" ，;")


def plain_notes(text: str | None) -> str | None:
    if not text:
        return None
    t = text
    for junk in ["phase7-batch-d", "phase5-batch-c", "global HRB-WY", "ExportStageDiscountApplierTest", "pathOverride 仍待 S3"]:
        t = t.replace(junk, "").strip()
    t = t.replace("；", "；").strip(" ；")
    mapping = {
        "分科室汇总": "账单可按科室出分科室汇总表",
        "结款合并组见": "与市五院主院区合并出结款函",
        "导出列裁剪已绑定": "账单里已去掉多余列",
        "3把16.5→8.91 特例见": "3 把 16.5 元特例：导出时改为 8.91 元",
        "≥3把3元 仅计价记录，非导出折扣": "满 3 把收 3 元（平时算价用，不是导出打折）",
    }
    for k, v in mapping.items():
        if k in t:
            return v
    return t if t else None


def split_rules(customer: dict[str, Any]) -> tuple[list[str], list[str]]:
    """返回 (仅导出时生效的规则说明, 平时算价规则说明)。"""
    export_lines: list[str] = []
    pricing_lines: list[str] = []
    for rule in customer.get("productRules") or []:
        if not rule.get("isActive", True):
            continue
        cj = rule.get("conditionsJson") or ""
        line = plain_rule_line(rule)
        if "exportApply" in cj:
            export_lines.append(line)
        else:
            pricing_lines.append(line)
    return export_lines, pricing_lines


HYBRID_STANDARD_PRICING_LINES = [
    "标准价·敷料包(无纺布)：W50/W60/W70→25元，W90→30元，W120/W150→35元",
    "驱血带（包名含「驱血带」）按敷料包(无纺布)同名 W 码分档计价",
]


def build_hospital_records(
    catalog: list[dict], customers: dict[str, Any], templates: dict[str, list[dict]]
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for entry in catalog:
        code = entry["code"]
        name = entry["folder"]
        customer = customers.get(code, {})
        export_only, pricing = split_rules(customer)
        if (customer.get("billingPricingMode") or "") == "hybrid":
            pricing = list(pricing) + HYBRID_STANDARD_PRICING_LINES
        tpl_lines: list[str] = []
        for t in templates.get(code, []):
            tpl_lines.extend(describe_template(t))

        records.append(
            {
                "code": code,
                "name": name,
                "export_types": EXPORT_FILE_TYPES.get(code, ["账单", "结款函"]),
                "bill_strategy": BILL_STRATEGY_PLAIN.get(
                    entry.get("billStrategy", ""), entry.get("billStrategy", "标准账单")
                ),
                "settlement_strategy": "标准结款函"
                if entry.get("settlementStrategy") == "standard_settlement"
                else entry.get("settlementStrategy", ""),
                "settlement_discount": plain_discount(entry.get("settlementDiscount")),
                "export_discount": plain_discount(entry.get("exportStageDiscount")),
                "notes": plain_notes(entry.get("notes")),
                "skip_reason": entry.get("skipReason"),
                "template_lines": tpl_lines,
                "export_only_rules": export_only,
                "pricing_rules": pricing,
                "pricing_mode": customer.get("billingPricingMode") or "standard",
                "active_rule_count": sum(
                    1 for r in (customer.get("productRules") or []) if r.get("isActive", True)
                ),
            }
        )
    return records


def build_docx(records: list[dict[str, Any]]) -> Document:
    doc = Document()
    set_doc_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("一期优先医院特色导出规则清单（请核对）")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"整理日期：{datetime.now().strftime('%Y年%m月%d日')}\n").font.size = Pt(9)
    meta.add_run("共 37 家医院 · 数据来自系统当前导出配置与计价规则").font.size = Pt(9)
    doc.add_paragraph()

    add_title(doc, "这份文件是干什么的？", 1)
    doc.add_paragraph(
        "这是一期优先上线的 37 家医院，在「导出账单、结款函、汇总表」时，程序里已经写好的特殊设置。"
        "请贵司按医院逐项核对：能不能导出哪些表、账单怎么排版、结款有没有打折、"
        "哪些包名在导出时有单独价钱。"
    )
    doc.add_paragraph(
        "说明：「平时算价规则」是导入原始表时就算好的价钱；"
        "「仅导出账单时生效」的规则是导出 Excel 时才改显示价钱，不影响系统里已算好的对账结果。"
    )

    add_title(doc, "汇总表", 1)
    summary_rows = []
    for rec in records:
        special = []
        if rec["export_discount"]:
            special.append("导出打折")
        if rec["settlement_discount"]:
            special.append("结款打折")
        if rec["export_only_rules"]:
            special.append(f"导出专价{len(rec['export_only_rules'])}条")
        if any("分科室" in x for x in rec.get("template_lines", [])):
            special.append("分科室账单")
        summary_rows.append(
            [
                rec["name"],
                "、".join(rec["export_types"]),
                "；".join(special) if special else "常规",
            ]
        )
    add_table(doc, ["医院", "可导出的表", "特色要点"], summary_rows)

    for idx, rec in enumerate(records, 1):
        add_title(doc, f"{idx}. {rec['name']}", 1)

        add_title(doc, "能导出哪些表", 2)
        for t in rec["export_types"]:
            doc.add_paragraph(t, style="List Bullet")

        add_title(doc, "账单与结款函格式", 2)
        rows = [
            ["账单", rec["bill_strategy"]],
            ["结款函", rec["settlement_strategy"]],
        ]
        if rec["settlement_discount"]:
            rows.append(["结款时额外处理", rec["settlement_discount"]])
        if rec["export_discount"]:
            rows.append(["导出账单时再打折", rec["export_discount"]])
        add_table(doc, ["项目", "说明"], rows)

        if rec["template_lines"]:
            add_title(doc, "已绑定的导出模板（排版/列）", 2)
            for line in rec["template_lines"]:
                doc.add_paragraph(line, style="List Bullet")

        if rec["export_only_rules"]:
            add_title(doc, "仅导出账单时生效的价钱（请重点核对）", 2)
            for line in rec["export_only_rules"]:
                doc.add_paragraph(line, style="List Bullet")
        else:
            doc.add_paragraph("该院没有单独写「仅导出时改价」的规则。")

        if rec["pricing_rules"]:
            add_title(doc, f"平时算价规则（共 {len(rec['pricing_rules'])} 条，导出账单也用这些价）", 2)
            show = rec["pricing_rules"][:15]
            for line in show:
                doc.add_paragraph(line, style="List Bullet")
            if len(rec["pricing_rules"]) > 15:
                doc.add_paragraph(f"……另有 {len(rec['pricing_rules']) - 15} 条，详见系统《医院特色计价规则清单》。")

        if rec["notes"]:
            doc.add_paragraph(f"备注：{rec['notes']}")
        if rec["skip_reason"]:
            doc.add_paragraph(f"待补材料：{rec['skip_reason']}")

        doc.add_paragraph("☐ 本院导出设置核对无误　☐ 有修改意见：________________")
        doc.add_paragraph("")

    add_title(doc, "客户确认回执", 1)
    doc.add_paragraph("联系人：____________　电话：____________　日期：____________")
    doc.add_paragraph("签字：____________")
    return doc


def build_md(records: list[dict[str, Any]]) -> str:
    lines = [
        "# 一期优先医院特色导出规则清单（请核对）",
        "",
        f"> 整理日期：{datetime.now().strftime('%Y-%m-%d')} · 共 37 家医院",
        "",
        "---",
        "",
    ]
    for idx, rec in enumerate(records, 1):
        lines.append(f"## {idx}. {rec['name']}")
        lines.append("")
        lines.append("**能导出：** " + "、".join(rec["export_types"]))
        lines.append("")
        lines.append(f"- 账单：{rec['bill_strategy']}")
        lines.append(f"- 结款函：{rec['settlement_strategy']}")
        if rec["settlement_discount"]:
            lines.append(f"- 结款额外：{rec['settlement_discount']}")
        if rec["export_discount"]:
            lines.append(f"- 导出打折：{rec['export_discount']}")
        if rec["template_lines"]:
            lines.append("")
            lines.append("**导出模板：**")
            for t in rec["template_lines"]:
                lines.append(f"- {t}")
        if rec["export_only_rules"]:
            lines.append("")
            lines.append("**仅导出时生效：**")
            for r in rec["export_only_rules"]:
                lines.append(f"- {r}")
        if rec["pricing_rules"]:
            lines.append("")
            lines.append(f"**平时算价（前 10 条）：**")
            for r in rec["pricing_rules"][:10]:
                lines.append(f"- {r}")
        if rec["notes"]:
            lines.append("")
            lines.append(f"备注：{rec['notes']}")
        lines.append("")
        lines.append("☐ 核对无误　☐ 有意见：________________")
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


def main() -> None:
    manifest = load_json(MANIFEST_PATH)
    export_seed = load_json(EXPORT_SEED)
    dept_seed = load_json(DEPT_SEED)
    templates = merge_templates(export_seed, dept_seed)
    catalog = export_seed["exportCatalog"]
    customers = manifest["customers"]
    records = build_hospital_records(catalog, customers, templates)

    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(build_md(records), encoding="utf-8")
    print(f"wrote {OUT_MD}")

    doc = build_docx(records)
    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
