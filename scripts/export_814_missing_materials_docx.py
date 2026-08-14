#!/usr/bin/env python3
"""Export concise missing-materials checklist for 814 batch (customer-facing Word)."""

from __future__ import annotations

import argparse
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TEST_CASE_DIR = ROOT / "测试用例"
OUT = TEST_CASE_DIR / "814新增缺材料清单-20260814.docx"

sys.path.insert(0, str(ROOT / "scripts"))
from batch_june_price_reconciliation import pick_month_pair  # noqa: E402
from ingest_bokang_814_batch import (  # noqa: E402
    FOUR_HOSPITAL_SPLIT,
    RAW_FILES,
    SPECIAL_PROC,
    STANDARD_PROC,
    STRICT_JULY_FOLDERS,
    load_manifest,
)
from special_v8_strict_excel_audit import V8_HOSPITALS  # noqa: E402

JULY_RAW_LABEL = "7月原始账单（系统导入格式）"
JULY_PROC_LABEL = "7月处理后账单"

JULY_SCOPE: set[str] = set()
for _, hospital, _, _, _ in RAW_FILES:
    JULY_SCOPE.add(hospital)
for _, hospital, _ in STANDARD_PROC:
    JULY_SCOPE.add(hospital)
for _, hospital, _ in FOUR_HOSPITAL_SPLIT:
    JULY_SCOPE.add(hospital)
JULY_SCOPE.update(
    {
        "博尚医院",
        "道里区妇幼保健院",
        "春语医美",
        "总工会",
        "基准生物",
        "索菲医美",
        "黑龙江省妇幼保健院（人口）",
    }
)


def month_prefixed_files(hospital_dir: Path, kind: str, month: int) -> list[Path]:
    sub = "原始表格" if kind == "raw" else "处理后表格"
    folder = hospital_dir / sub
    if not folder.is_dir():
        return []
    prefix = f"{month}月__"
    return sorted(p for p in folder.glob("*.xlsx") if p.name.startswith(prefix))


def july_gaps(hospital_dir: Path) -> list[str]:
    if not hospital_dir.is_dir():
        return [JULY_RAW_LABEL, JULY_PROC_LABEL]

    raw7 = month_prefixed_files(hospital_dir, "raw", 7)
    proc7 = month_prefixed_files(hospital_dir, "proc", 7)
    gaps: list[str] = []

    if raw7 and not proc7:
        gaps.append(JULY_PROC_LABEL)
    elif proc7 and not raw7:
        gaps.append(JULY_RAW_LABEL)
    elif not raw7 and not proc7:
        raw, proc, _ = pick_month_pair(hospital_dir, 7)
        if raw and proc:
            return []
        if raw and not proc:
            gaps.append(JULY_PROC_LABEL)
        elif proc and not raw:
            gaps.append(JULY_RAW_LABEL)
        else:
            gaps.extend([JULY_RAW_LABEL, JULY_PROC_LABEL])
    return gaps


def extra_month_gaps(hospital: str, hospital_dir: Path) -> list[str]:
    gaps: list[str] = []

    if hospital == "黑龙江九洲妇科医院":
        if month_prefixed_files(hospital_dir, "proc", 3) and not pick_month_pair(hospital_dir, 3)[0]:
            gaps.append("3月原始账单（系统导入格式）")

    elif hospital == "哈尔滨工程大学医院":
        if month_prefixed_files(hospital_dir, "proc", 6) and not month_prefixed_files(hospital_dir, "raw", 6):
            raw6, proc6, _ = pick_month_pair(hospital_dir, 6)
            if proc6 and (not raw6 or not month_prefixed_files(hospital_dir, "raw", 6)):
                gaps.append("6月原始账单（系统导入格式）")

    elif hospital == "松电慢病":
        if not month_prefixed_files(hospital_dir, "raw", 3):
            gaps.append("3月原始账单（系统导入格式）")

    elif hospital == "黑龙江省海员总医院（松北）":
        if not month_prefixed_files(hospital_dir, "raw", 3) and not month_prefixed_files(hospital_dir, "raw", 4):
            gaps.append("3–4月原始账单（系统导入格式）")

    elif hospital == "航天风华":
        if month_prefixed_files(hospital_dir, "proc", 8) and not month_prefixed_files(hospital_dir, "raw", 8):
            gaps.append("8月原始账单（系统导入格式）")

    elif hospital == "哈尔滨市第五医院（二门诊）":
        raw6, proc6, _ = pick_month_pair(hospital_dir, 6)
        if not raw6 or not proc6:
            gaps.append("6月原始账单（系统导入格式）")

    elif hospital == "省监狱管理局":
        gaps.extend(["原始账单", "处理后账单"])

    return gaps


def scan_scope() -> list[str]:
    hospitals: set[str] = set()
    for _, hospital, _, _, _ in RAW_FILES:
        hospitals.add(hospital)
    for _, hospital, _ in STANDARD_PROC:
        hospitals.add(hospital)
    for _, hospital, _, _ in SPECIAL_PROC:
        hospitals.add(hospital)
    for _, hospital, _ in FOUR_HOSPITAL_SPLIT:
        hospitals.add(hospital)
    for entry in load_manifest().get("entries", []):
        hospital = entry.get("hospital")
        if hospital and hospital != "待匹配":
            hospitals.add(hospital)
    for h in V8_HOSPITALS:
        if h.folder:
            hospitals.add(h.folder)
    hospitals.add("省监狱管理局")
    hospitals -= set(STRICT_JULY_FOLDERS)
    return sorted(hospitals)


def collect_missing() -> list[tuple[str, list[str]]]:
    rows: list[tuple[str, list[str]]] = []
    for hospital in scan_scope():
        hospital_dir = TEST_CASE_DIR / hospital
        items: list[str] = []
        if hospital == "省监狱管理局":
            items.extend(extra_month_gaps(hospital, hospital_dir))
        else:
            if hospital in JULY_SCOPE:
                items.extend(july_gaps(hospital_dir))
            items.extend(extra_month_gaps(hospital, hospital_dir))
        deduped: list[str] = []
        seen: set[str] = set()
        for x in items:
            if x not in seen:
                seen.add(x)
                deduped.append(x)
        if deduped:
            rows.append((hospital, deduped))
    return rows


def set_doc_font(doc: Document, name: str = "微软雅黑") -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def build_document(missing_rows: list[tuple[str, list[str]]]) -> Document:
    doc = Document()
    set_doc_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("814 新增账单 — 缺材料清单")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(date.today().isoformat()).font.size = Pt(10)

    for hospital, items in missing_rows:
        h = doc.add_paragraph()
        hr = h.add_run(hospital)
        hr.bold = True
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    return doc


def main() -> int:
    p = argparse.ArgumentParser(description="导出 814 缺材料清单 Word")
    p.add_argument("--write", action="store_true", help="写入 docx")
    p.add_argument("--out", type=Path, default=OUT)
    args = p.parse_args()
    missing_rows = collect_missing()
    doc = build_document(missing_rows)
    if args.write:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(args.out)
        print(f"已写入: {args.out} ({len(missing_rows)} 家)")
    else:
        print(f"预览 {len(missing_rows)} 家，使用 --write 写入 {args.out}")
        for hospital, items in missing_rows[:8]:
            print(f"  {hospital}: {', '.join(items)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
