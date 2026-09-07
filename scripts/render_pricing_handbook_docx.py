#!/usr/bin/env python3
"""将 docs/通用计价规则手册.md 渲染为面向客户的 Word 文档（正式商务排版）。

可重复运行：手册 md 更新后重新执行本脚本即可再生成 docx，脚本不修改 md 本身。

用法：
    python3 scripts/render_pricing_handbook_docx.py

支持的 markdown 子集：# 标题层级、| 表格 |（含 \\| 转义与 |---| 分隔行）、
- 无序列表、1. 有序列表、> 引用块、**加粗**、`行内代码`（按普通文本渲染）、--- 分隔线。
生成后自动重新打开 docx 做校验（表格数与 md 一致、无 ** / |--- 等解析残留）。
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs/通用计价规则手册.md"
OUT_DOCX = ROOT / "docs/通用计价规则手册.docx"

DOC_TITLE = "通用计价规则手册"
DOC_SUBTITLE = "系统默认计价方式说明（未约定特殊价格的医院均按此执行）"
DOC_DATE_FALLBACK = "2026-09-07"

BODY_EA_FONT = "宋体"
BODY_ASCII_FONT = "Times New Roman"
HEAD_EA_FONT = "微软雅黑"
HEAD_ASCII_FONT = "微软雅黑"

# md 标题级别 -> Word 标题级别/字号：## 章 -> 16pt，### 节 -> 14pt，#### -> 12pt
HEADING_STYLE = {
    2: ("Heading 1", 16, RGBColor(0x1F, 0x4E, 0x79)),
    3: ("Heading 2", 14, RGBColor(0x1F, 0x4E, 0x79)),
    4: ("Heading 3", 12, RGBColor(0x40, 0x40, 0x40)),
}

PAGE_WIDTH_CM = 21.0
MARGIN_LR_CM = 2.2
TABLE_BODY_PT = 9.5
HEADER_FILL = "D9E2F3"
QUOTE_FILL = "F2F2F2"

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|?\s*$", re.M)


# ---------------------------------------------------------------- 基础工具

def set_run_font(run, *, ea=BODY_EA_FONT, ascii_=BODY_ASCII_FONT,
                 size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = ascii_
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, *, ea, ascii_, size=None, bold=None, color=None) -> None:
    style.font.name = ascii_
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_)
    rfonts.set(qn("w:hAnsi"), ascii_)
    rfonts.set(qn("w:eastAsia"), ea)


def add_inline_runs(p, text, *, size=None, base_bold=False, color=None,
                    ea=BODY_EA_FONT, ascii_=BODY_ASCII_FONT) -> None:
    """把含 **加粗** / `代码` 的文本写入段落，加粗真正生效，代码按普通文本渲染。"""
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        bold = base_bold
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            tok = tok[2:-2]
            bold = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            tok = tok[1:-1]
        run = p.add_run(tok)
        set_run_font(run, ea=ea, ascii_=ascii_, size=size, bold=bold, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def shade_paragraph(p, fill: str) -> None:
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def paragraph_border(p, edge: str, color: str, sz: str) -> None:
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), color)
    p_bdr.append(el)


def display_len(text: str) -> int:
    return sum(2 if ord(ch) > 0x2E7F else 1 for ch in text)


# ---------------------------------------------------------------- markdown 解析

def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|") and not line.endswith("\\|"):
        line = line[:-1]
    parts = re.split(r"(?<!\\)\|", line)
    return [p.replace("\\|", "|").strip() for p in parts]


def is_sep_cells(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def parse_blocks(lines: list[str]):
    """把 md 行序列解析为块列表：(kind, payload)。"""
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
        elif s.startswith(">"):
            qlines = []
            while i < n and lines[i].strip().startswith(">"):
                qlines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(("quote", [q for q in qlines if q]))
        elif s.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            rows = [split_table_row(x) for x in tbl]
            header = rows[0]
            body = [r for r in rows[1:] if not is_sep_cells(r)]
            ncol = len(header)
            body = [(r + [""] * ncol)[:ncol] for r in body]
            blocks.append(("table", (header, body)))
        elif re.fullmatch(r"-{3,}", s):
            blocks.append(("hr", None))
            i += 1
        elif re.match(r"^\s*[-*]\s+", raw):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                line = lines[i]
                indent = len(line) - len(line.lstrip())
                items.append((indent // 2, re.sub(r"^\s*[-*]\s+", "", line).strip()))
                i += 1
            blocks.append(("ul", items))
        elif re.match(r"^\d+[.、]\s+", s):
            items = []
            while i < n and re.match(r"^\s*\d+[.、]\s+", lines[i].strip()):
                items.append(lines[i].strip())
                i += 1
            blocks.append(("ol", items))
        else:
            blocks.append(("para", s))
            i += 1
    return blocks


# ---------------------------------------------------------------- docx 渲染

def setup_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, ea=BODY_EA_FONT, ascii_=BODY_ASCII_FONT, size=Pt(10.5))
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(4)

    for lvl, (style_name, size, color) in HEADING_STYLE.items():
        st = doc.styles[style_name]
        set_style_font(st, ea=HEAD_EA_FONT, ascii_=HEAD_ASCII_FONT,
                       size=Pt(size), bold=True, color=color)
        st.paragraph_format.space_before = Pt(14 if lvl == 2 else 10)
        st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width = Cm(PAGE_WIDTH_CM)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(MARGIN_LR_CM)
    sec.right_margin = Cm(MARGIN_LR_CM)
    sec.different_first_page_header_footer = True  # 封面不显示页眉页脚

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(DOC_TITLE)
    set_run_font(run, size=Pt(9), color=RGBColor(0x7F, 0x7F, 0x7F))
    paragraph_border(hp, "bottom", "BFBFBF", "4")

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("第 ")
    set_run_font(r1, size=Pt(9), color=RGBColor(0x7F, 0x7F, 0x7F))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fr = OxmlElement("w:r")
    fr_pr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    fr_pr.append(sz)
    ft = OxmlElement("w:t")
    ft.text = "1"
    fr.append(fr_pr)
    fr.append(ft)
    fld.append(fr)
    fp._p.append(fld)
    r2 = fp.add_run(" 页")
    set_run_font(r2, size=Pt(9), color=RGBColor(0x7F, 0x7F, 0x7F))


def add_cover(doc: Document, doc_date: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DOC_TITLE)
    set_run_font(run, ea=HEAD_EA_FONT, ascii_=HEAD_ASCII_FONT,
                 size=Pt(26), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_border(line, "bottom", "1F4E79", "8")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DOC_SUBTITLE)
    set_run_font(run, size=Pt(13), color=RGBColor(0x40, 0x40, 0x40))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"日期：{doc_date}")
    set_run_font(run, size=Pt(11), color=RGBColor(0x7F, 0x7F, 0x7F))
    doc.add_page_break()


def add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    paragraph_border(p, "bottom", "BFBFBF", "6")


def add_quote(doc: Document, qlines: list[str]) -> None:
    for q in qlines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        shade_paragraph(p, QUOTE_FILL)
        paragraph_border(p, "left", "A6A6A6", "18")
        add_inline_runs(p, q, size=Pt(10), color=RGBColor(0x40, 0x40, 0x40))


def column_widths(header: list[str], rows: list[list[str]]) -> list[float]:
    usable = PAGE_WIDTH_CM - 2 * MARGIN_LR_CM
    ncol = len(header)
    weights = []
    for c in range(ncol):
        longest = display_len(header[c])
        for r in rows:
            longest = max(longest, display_len(r[c]))
        weights.append(min(max(longest, 4), 24))
    total = sum(weights)
    return [usable * w / total for w in weights]


def add_md_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    ncol = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = column_widths(header, rows)

    for c, h in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.width = Cm(widths[c])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        add_inline_runs(p, h, size=Pt(TABLE_BODY_PT), base_bold=True)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.width = Cm(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            add_inline_runs(p, val, size=Pt(TABLE_BODY_PT))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    for run in spacer.runs:
        run.font.size = Pt(4)


def render(doc: Document, blocks) -> None:
    body_started = False
    for kind, payload in blocks:
        # 封面区：md 的一级标题、紧随其后的引用块与分隔线已由封面页承载，跳过
        if not body_started and kind in ("heading", "quote", "hr"):
            if kind == "heading" and payload[0] != 1:
                body_started = True
            else:
                continue
        if not body_started:
            body_started = True

        if kind == "heading":
            level, text = payload
            style_name = HEADING_STYLE.get(level, HEADING_STYLE[4])[0]
            p = doc.add_paragraph(style=style_name)
            add_inline_runs(p, text, ea=HEAD_EA_FONT, ascii_=HEAD_ASCII_FONT)
        elif kind == "quote":
            add_quote(doc, payload)
        elif kind == "table":
            add_md_table(doc, payload[0], payload[1])
        elif kind == "hr":
            add_hr(doc)
        elif kind == "ul":
            for indent, text in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.75 + 0.5 * indent)
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, text)
        elif kind == "ol":
            for text in payload:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75)
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, text)
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, payload)


# ---------------------------------------------------------------- 校验

def validate(md_text: str, doc: Document) -> None:
    md_tables = len(TABLE_SEP_RE.findall(md_text))
    doc_tables = len(doc.tables)
    print(f"[校验] md 表格数 = {md_tables}，docx 表格数 = {doc_tables} -> "
          + ("一致" if md_tables == doc_tables else "不一致！"))

    print(f"[校验] docx 正文段落数 = {len(doc.paragraphs)}")

    residue = []
    texts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    for t in texts:
        for marker in ("**", "`", "|--", "---"):
            if marker in t:
                residue.append((marker, t[:60]))
    if residue:
        print(f"[校验] 发现 {len(residue)} 处解析残留：")
        for marker, snippet in residue[:10]:
            print(f"    残留 {marker!r}: {snippet}")
    else:
        print("[校验] 无解析残留（未发现 **、`、|--- 等）")

    for t in doc.tables:
        first = t.rows[0].cells[0].text.strip()
        if first == "袋宽":
            print("[抽查] 高温纸塑袋袋费表内容：")
            for row in t.rows:
                print("    | " + " | ".join(c.text.strip() for c in row.cells) + " |")
            cells = [c.text.strip() for c in t.rows[1].cells]
            ok = cells == ["袋费（元）", "2.5", "5.5", "7.5", "10.5"]
            print("[抽查] 袋费数值行 " + ("正确" if ok else "异常！"))
            break


# ---------------------------------------------------------------- 主流程

def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    m = re.search(r"日期：(\d{4}-\d{2}-\d{2})", md_text)
    doc_date = m.group(1) if m else DOC_DATE_FALLBACK

    blocks = parse_blocks(md_text.splitlines())

    doc = Document()
    setup_document(doc)
    add_cover(doc, doc_date)
    render(doc, blocks)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"已生成 {OUT_DOCX}")

    check = Document(OUT_DOCX)
    validate(md_text, check)


if __name__ == "__main__":
    sys.exit(main())
