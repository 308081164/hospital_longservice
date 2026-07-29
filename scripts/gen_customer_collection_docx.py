#!/usr/bin/env python3
"""生成「测试材料补充清单」Word 文档（客户阅读版）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "测试用例" / "特色账单测试材料补充清单（客户版）.docx"

FONT = "微软雅黑"
COLOR_TITLE = RGBColor(0x1A, 0x47, 0x7A)
COLOR_H1 = RGBColor(0xC0, 0x39, 0x2B)
COLOR_H2 = RGBColor(0x1A, 0x47, 0x7A)
COLOR_H3 = RGBColor(0x2E, 0x86, 0xAB)


def set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run(run, size=size)


def set_cell_shading(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(hdr[j], "1A477A")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run(run, size=10)
            if i % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return table


def add_hospital_block(doc, name: str, status: str, need_title: str, items: list[str]):
    add_para(doc, name, size=12, bold=True, color=COLOR_H3, space_after=4)
    add_para(doc, f"目前情况：{status}", bold=True, space_after=4)
    add_para(doc, need_title, bold=True, space_after=2)
    add_bullets(doc, items, size=10)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    add_para(doc, "特色账单系统", size=22, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "测试材料补充清单", size=18, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "（供客户查阅）", size=12, color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "更新日期：2026 年 7 月 29 日", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # 说明
    add_para(doc, "这份清单是做什么的？", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_para(
        doc,
        "我们在核对「特色账单系统」算出来的金额，是否和贵司平时手工处理后的结果一致。"
        "核对时，每家医院、每个账期都需要两份 Excel：",
        space_after=6,
    )
    add_bullets(doc, [
        "第一份：系统计价前导出的原始账单（还没做特色调价）",
        "第二份：贵司确认过的正确账单（或结款函）",
    ])
    add_para(
        doc,
        "两份必须是同一账期（起止日期一致），并且能按发货单号、包名、包数一行一行对上。"
        "做了特色调价的行，处理后价格应和原始不同；没做调价的行，两边价格应相同。",
        space_after=6,
    )
    add_para(
        doc,
        "目前多数医院 6 月已核对完毕。下面列出还需要补充材料的医院，按紧急程度排列。"
        "收到一批材料，我们会核对一批并反馈结果。",
        space_after=14,
    )

    # 进度概况
    add_para(doc, "目前整体进度（简要）", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_bullets(doc, [
        "共 37 家医院纳入本轮核对",
        "其中 28 家：6 月账单和结款函都已核对通过",
        "4 家：因缺少原始明细或成套包拆分说明，6 月账单暂时对不上",
        "4 家：6 月账单已通过，但还缺结款函参考表",
        "哈尔滨工程大学医院：5 月原始账单已于 7 月 29 日收到，正在安排核对",
    ])
    doc.add_page_break()

    # 统一要求
    add_para(doc, "提供材料时请注意", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_bullets(doc, [
        "同一账期请成对提供：原始账单 + 处理后账单（或结款函）",
        "账期起止日期必须一致",
        "文件格式优先 Excel；结款函如果只有 Word，也请尽量转成 Excel",
        "文件名建议写上医院全称和月份，方便归档",
    ])
    doc.add_paragraph()

    # 第一节
    add_para(doc, "一、最优先：缺材料导致 6 月账单对不上（共 4 家）", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(
        doc,
        "以下 4 家医院，系统算出的 6 月账单和贵司确认版对不上，主要原因是原始表里缺明细，"
        "或成套器械包在贵司账单里拆成了多行、但系统这边没有对应规则。"
        "请优先补充下列材料。",
        space_after=10,
    )

    add_hospital_block(
        doc,
        "1. 国药总医院主院区",
        "6 月账单金额差约 696 元。系统导出有 131 行，贵司确认版有 206 行，中间缺了成套包的明细行。",
        "请提供以下任一项（或多项）：",
        [
            "成套器械包拆分说明：例如「产包」「腔镜包」在贵司账单里拆成了哪些独立明细行",
            "或确认：汽轮机 6 月原始账单里，是否应已包含全部组件明细",
            "若同一包名两边数量不一致（如治疗盘器械、外科缝合包），请说明应以哪边为准",
            "对应账期：5 月 26 日 — 6 月 25 日",
        ],
    )

    add_hospital_block(
        doc,
        "2. 国药总医院第二院区",
        "与主院区类似，6 月账单差约 121.5 元。系统 43 行，贵司确认版 65 行。",
        "请提供：",
        [
            "同主院区的成套包拆分说明",
            "或确认：电机厂 6 月原始账单是否缺组件明细",
            "对应账期：5 月 26 日 — 6 月 25 日",
        ],
    )

    add_hospital_block(
        doc,
        "3. 哈尔滨市第二医院",
        "6 月账单差约 1.19 万元。贵司确认版比现有原始账单多了 7 个供应商的明细页，"
        "这些明细在原始账单里找不到。",
        "请提供：",
        [
            "2026 年 6 月外来器械/供应商补录明细（Excel）",
            "需包含以下供应商（或等价明细）：上海尔欢、大博、尔欢、捷迈得、星檀、纽枫、钇嵩 等",
            "与 6 月市二院确认版账单同一账期",
        ],
    )

    add_hospital_block(
        doc,
        "4. 黑龙江省第二医院（松北院区）",
        "6 月账单差约 8,743 元。一部分是原始账单里缺外来器械、钉盒等明细；"
        "另一部分是成套包（如宫腔镜包）在贵司账单里拆行、系统这边对不上。",
        "请提供：",
        [
            "外来器械、钉盒等补录明细（贵司确认版有、原始账单里没有的约 110 行）",
            "成套包拆分说明（如宫腔镜包拆成哪些独立行）",
            "若「盆1碗2盘2/W9050」「上肢钉盒（三）」等应在原始账单出现，请确认是否需重新导出",
        ],
    )
    doc.add_page_break()

    # 第二节 结款函
    add_para(doc, "二、次优先：6 月账单已对，还缺结款函（共 4 家）", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(
        doc,
        "以下医院 6 月账单已核对通过，但还缺少贵司确认过的结款函 Excel，暂时无法核对结款函。"
        "（第一节 4 家医院的结款函，需等账单材料补齐后再核对。）",
        space_after=8,
    )
    add_table(
        doc,
        ["医院", "还缺什么"],
        [
            ["国药总医院第三院区", "6 月同账期的结款函（Excel）"],
            ["香坊中医院", "6 月同账期的结款函（Excel）"],
            ["哈尔滨长健医院", "6 月同账期的结款函（Excel）"],
            ["哈尔滨市第五医院（二门诊）", "6 月同账期的结款函（Excel）"],
        ],
        [5.5, 10],
    )

    # 第三节 历史账期
    add_para(doc, "三、补充核对用：其它月份的材料", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(
        doc,
        "以下不影响 6 月主核对，但有助于把特色计价规则测得更全。"
        "部分医院 6 月原始和处理后价格完全一样，说明 6 月可能没有改价，"
        "需要再找一个「确实有改价」的月份来验证。",
        space_after=8,
    )

    add_hospital_block(
        doc,
        "5. 哈尔滨市红十字妇产医院",
        "6 月已核对；4 月、5 月还缺贵司确认版账单，暂时测不了。",
        "请提供：",
        [
            "4 月确认版账单（4 月原始账单已有）",
            "5 月确认版账单（5 月原始账单已有）",
            "建议 4、5 月尽量包含低温、湿化瓶、T 型管、喉镜/软管等有改价的明细",
        ],
    )

    add_hospital_block(
        doc,
        "6. 哈尔滨工业大学医院",
        "4 月、5 月已核对；6 月原始账单和处理后账单不是同一账期，对不上。",
        "请提供：",
        [
            "2026 年 6 月 15 日 — 7 月 14 日的原始账单（确认版已有）",
            "建议包含口腔相关明细（针类、洁牙尖、成型片、克氏针、车针等）",
        ],
    )

    add_para(doc, "7. 6 月两边价格一样，建议补「有改价月份」的 10 家", size=12, bold=True, color=COLOR_H3, space_after=4)
    add_para(
        doc,
        "这 10 家医院 6 月原始和处理后价格相同，只能证明「没算错」，不能证明特色规则算对了。"
        "请在下列月份补原始 + 确认版成对账单：",
        space_after=6,
    )
    add_table(
        doc,
        ["序号", "医院", "建议补哪个月", "说明"],
        [
            ["1", "太平人民医院", "5 月（最优先）", "约 101 处价格有差别"],
            ["2", "南岗区妇产医院", "5 月", "约 10 处"],
            ["3", "国药总医院第二院区", "2 月", "约 3 处"],
            ["4", "祖研-黑龙江省中医医院（香安院区）", "4 月", "约 3 处"],
            ["5", "黑龙江维多利亚妇产医院", "5 月", "约 2 处"],
            ["6", "呼兰区红十字医院", "5 月", "约 2 处"],
            ["7", "国药总医院主院区", "3 月", "约 1 处"],
            ["8", "黑龙江省社会康复医院", "4 月", "约 1 处"],
            ["9", "呼兰中医院", "5 月", "约 1 处"],
            ["10", "黑龙江中医药大学附属第二医院（哈南分院）", "5 月", "约 1 处"],
        ],
        [1.2, 6.5, 2.8, 4],
    )
    add_para(
        doc,
        "另外：黑龙江省第二医院（南岗院区）6 月也无差别；若方便，请一并提供 4 月原始 + 确认版（约 3 处口腔科成型夹相关差异）。",
        size=10,
        space_after=14,
    )
    doc.add_page_break()

    # 第四节 早期账期
    add_para(doc, "四、可选：更早月份的材料（不着急）", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_para(doc, "若方便，下列医院可补 1–3 月原始 + 确认版成对账单，用于更早月份抽查：", space_after=4)
    add_bullets(doc, [
        "黑龙江中医药大学附属第三医院（电力）— 1 月",
        "道外区人民医院 — 1 月",
        "黑龙江维多利亚妇产医院 — 1 月",
        "黑龙江九洲妇科医院 — 1 月",
        "哈尔滨仁胜医院 — 1 月",
        "哈尔滨冰城医疗美容医院 — 1 月",
        "哈尔滨华夏眼科医院 — 1 月",
        "黑龙江省医院（南岗/香坊）— 1 月",
        "悦美芳华医疗门诊医院 — 1 月",
        "黑龙江省第二医院（南岗院区）— 1 月",
        "哈尔滨市呼兰区第一人民医院 — 1 月",
    ], size=10)
    add_para(doc, "以下医院确认版已有，只缺原始账单：", bold=True, space_after=4)
    add_bullets(doc, [
        "黑龙江奥兰医院 — 1 月",
        "黑龙江省森工总医院 — 1 月",
        "黑龙江省森工总医院（香坊院区）— 1 月",
        "黑龙江省森工总医院（松北院区）— 1 月",
        "黑龙江省骨伤科医院 — 1 月",
        "黑龙江省骨伤科医院（二院区）— 1 月",
        "黑龙江省骨伤科医院（三院区）— 1 月",
        "黑龙江省骨伤科医院（四院区）— 1 月",
    ], size=10)
    doc.add_paragraph()

    # 已收材料
    add_para(doc, "五、已经收到、不用再补的", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_hospital_block(
        doc,
        "哈尔滨工程大学医院",
        "5 月原始账单已于 2026 年 7 月 29 日收到，与 5 月确认版账期一致（5 月 1 日 — 5 月 31 日）。",
        "说明：",
        [
            "原始账单和 5 月确认版已成对，正在安排系统核对",
            "此项无需再向贵司索要",
        ],
    )

    # 无需补件
    add_para(doc, "六、暂不需要补材料的情况", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_para(doc, "6 月已核对通过、暂不需新材料", bold=True, space_after=4)
    add_para(
        doc,
        "除上文所列外，其余多数医院 6 月账单和结款函都已核对通过，暂时不需要新材料。",
        space_after=8,
    )
    add_para(doc, "太平人民医院（金额差很小）", bold=True, space_after=4)
    add_para(
        doc,
        "6 月账单和结款函整体正确，只剩约 20 元的小数尾差，属于阶梯价四舍五入方式不同，不影响使用。"
        "若贵司有书面说明「按每行四舍五入」还是「按总额四舍五入」，可提供以便完全对齐；不提供也不影响验收。",
        space_after=8,
    )
    add_para(doc, "黑龙江省社会康复医院（账套说明）", bold=True, space_after=4)
    add_para(
        doc,
        "6 月核对以「省康复」账套为准；「监狱」账套是另一套独立账目，两者分开核对，无需额外补材料，"
        "只需确认以哪套为准即可。",
        space_after=14,
    )

    # 转发话术
    add_para(doc, "附：可转发的一段话", size=14, bold=True, color=COLOR_H2, space_after=6)
    quote = (
        "各位老师好，\n\n"
        "特色账单系统 6 月核对进展：37 家医院中，28 家账单和结款函都已通过。\n\n"
        "还需要补充材料的，按优先级如下：\n\n"
        "【最优先 · 4 家，6 月账单对不上】\n"
        "① 国药总院主院区、第二院区 — 成套器械包怎么拆成明细行的说明；\n"
        "② 市二院 — 6 月外来器械/供应商补录明细（7 个供应商）；\n"
        "③ 省二松北 — 外来器械补录 + 成套包拆分说明。\n\n"
        "【次优先 · 4 家，缺结款函】\n"
        "国药三院、香坊中医院、长健、市五二门诊 — 6 月同账期结款函 Excel。\n\n"
        "【补充核对用】\n"
        "红十字妇产 — 4 月、5 月确认版账单；\n"
        "哈工大 — 6 月 15 日至 7 月 14 日原始账单；\n"
        "另有 10 家建议补「有改价月份」的成对账单（详见正文表格）。\n\n"
        "同一账期请成对提供「原始 + 确认版」；有改价的行，确认版价格应和原始不同。"
        "材料可以分批给，收到一批我们核对一批。感谢配合！"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(quote)
    set_run(run, size=10, italic=True)

    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    build()
