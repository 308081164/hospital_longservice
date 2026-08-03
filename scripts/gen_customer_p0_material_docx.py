#!/usr/bin/env python3
"""生成「四家医院 6 月明细补充清单」Word 文档（客户阅读版 · 2026-08-01）。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "测试用例" / "四家医院6月账单明细补充清单（客户版）.docx"

FONT = "微软雅黑"
COLOR_TITLE = RGBColor(0x1A, 0x47, 0x7A)
COLOR_H1 = RGBColor(0xC0, 0x39, 0x2B)
COLOR_H2 = RGBColor(0x1A, 0x47, 0x7A)
COLOR_H3 = RGBColor(0x2E, 0x86, 0xAB)


def set_run(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run(run, size=size)


def set_cell_shading(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(hdr[j], "1A477A")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run(run, size=10)
            if i % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return table


def add_hospital_block(doc, name: str, status: str, need_title: str, items: list[str]):
    add_para(doc, name, size=12, bold=True, color=COLOR_H3, space_after=4)
    add_para(doc, f"目前情况：{status}", bold=True, space_after=4)
    add_para(doc, need_title, bold=True, space_after=2)
    add_bullets(doc, items, size=10)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    add_para(doc, "特色账单系统", size=22, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "测试材料补充清单", size=18, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "（四家医院 · 6 月账单明细补充说明 · 供客户查阅）", size=12, color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "更新日期：2026 年 8 月 1 日", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_para(doc, "这份清单是做什么的？", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_para(
        doc,
        "我们在核对「特色账单系统」导出的 6 月账单，是否和贵司平时手工处理后的结果一致。"
        "下面 4 家医院的 6 月原始账单 + 处理后账单我们均已收到，常规对账材料是齐的。",
        space_after=6,
    )
    add_para(
        doc,
        "但在逐行比对时，系统导出的账单仍比贵司确认版「少行、少金额」。"
        "排查后发现：缺的不是「再要一份总账单」，而是三类「拆分明细」——"
        "就像给了总账，还需要分录明细才能一行一行对上。",
        space_after=6,
    )
    add_para(doc, "收到下列材料后，我们会做单院复测；对齐一家、反馈一家。", space_after=14)

    add_para(doc, "目前整体进度（简要）", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_bullets(doc, [
        "共 37 家医院纳入本轮 6 月核对",
        "其中 33 家：6 月账单已核对通过",
        "以下 4 家：6 月总账单已有，但缺「拆包 / 分厂家 / 第三部分」明细，暂时对不齐",
        "上述 4 家的结款函，需等账单明细补齐后再核对",
    ])
    doc.add_paragraph()

    add_table(
        doc,
        ["医院", "目前差额（约）", "说明"],
        [
            ["国药总医院主院区", "696 元", "系统导出少 8 行"],
            ["国药总医院第二院区", "121 元", "系统导出少 11 行"],
            ["哈尔滨市第二医院", "11,900 元", "系统导出少 36 行"],
            ["黑龙江省第二医院（松北院区）", "8,743 元", "系统导出少 148 行"],
        ],
        [5.5, 3.2, 6.8],
    )
    doc.add_page_break()

    add_para(doc, "提供材料时请注意", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_bullets(doc, [
        "账期与现有 6 月验收批次一致（国药主/二为 5 月 26 日 — 6 月 25 日）",
        "文件格式优先 Excel（.xlsx）；Word 表格也可，但需能逐行对照",
        "文件名建议写上医院全称和月份，方便归档",
        "不必再发一遍已有的 6 月总账单 / 结款函，除非贵司确认原始版本有更新",
    ])
    doc.add_paragraph()

    add_para(doc, "一、国药总医院主院区 / 第二院区", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(doc, "需要补充：「整包拆零件」对照表", size=12, bold=True, color=COLOR_H3, space_after=6)
    add_para(doc, "贵司原始账单里往往是一行写「某某包 × 1」；处理后账单里会变成多行零件，例如：", space_after=4)
    add_bullets(doc, ["原始：产包 × 1", "处理后：持针器-1/… × 1、弯血管钳-2/… × 2 等多条"], size=10)
    add_para(
        doc,
        "我们系统导入时只看到了「整包」那一行，不知道一个包里面具体有哪些零件、各几件，"
        "所以导出的账单比贵司少行、少金额。",
        space_after=8,
    )

    add_hospital_block(
        doc,
        "1. 国药总医院主院区",
        "6 月账单差约 696 元。系统导出 131 行，贵司确认版 206 行。",
        "请提供以下任一项（或多项）：",
        [
            "拆包对照表（Excel）：原始包名 → 拆开后每一行的名称、数量、单价（如有）",
            "账期：5 月 26 日 — 6 月 25 日（与现有 6 月处理后表一致）",
            "或请确认：汽轮机 6 月原始账单里，是否本来就应该按「零件行」录入？",
            "若应录入但仍为整包，请补一份已按零件展开的原始明细",
            "若有「一个包算几件」的特殊算法（如 packCount），请给 1～2 个算例或文字说明",
        ],
    )

    add_para(doc, "拆包对照表示例（表头可参考，不必完全一致）：", bold=True, space_after=4)
    add_table(
        doc,
        ["原始包名", "零件名称", "数量", "备注"],
        [
            ["产包", "持针器-1/Z1026", "1", ""],
            ["产包", "弯血管钳-2/…", "2", ""],
            ["腔镜包", "（请按贵司实际填写）", "…", ""],
        ],
        [3.5, 5.5, 2, 4.5],
    )

    add_hospital_block(
        doc,
        "2. 国药总医院第二院区",
        "与主院区类似，6 月账单差约 121.5 元。系统 43 行，贵司确认版 65 行。",
        "请提供：",
        [
            "同主院区的拆包对照表（格式同上）",
            "或确认：电机厂 6 月原始账单是否缺组件明细",
            "账期：5 月 26 日 — 6 月 25 日",
        ],
    )
    doc.add_page_break()

    add_para(doc, "二、哈尔滨市第二医院", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(doc, "需要补充：供应商（外来器械）补录明细", size=12, bold=True, color=COLOR_H3, space_after=6)
    add_para(
        doc,
        "贵司 6 月处理后账单里，比我们的 6 月原始账单多了 7 个供应商/厂家的明细页，"
        "这些明细在现有「市二院 6 月账单」里完全没有，合计约 1.19 万元、36 行。",
        space_after=8,
    )
    add_hospital_block(
        doc,
        "3. 哈尔滨市第二医院",
        "6 月账单差约 11,900 元。",
        "请提供：",
        [
            "2026 年 6 月上述供应商/厂家的原始明细 Excel",
            "需包含以下供应商（或等价明细）：上海尔欢、大博、尔欢、捷迈得、星檀、纽枫、钇嵩 等",
            "可以是多个 sheet 在一个文件里，或与贵司 6 月确认版结构一致的补录表",
            "与贵司「6 月市二院确认版账单」同一账期、能一一对应即可",
        ],
    )
    doc.add_page_break()

    add_para(doc, "三、黑龙江省第二医院（松北院区）", size=14, bold=True, color=COLOR_H1, space_after=6)
    add_para(doc, "需要补充：第三部分明细 + 拆包对照", size=12, bold=True, color=COLOR_H3, space_after=6)
    add_para(doc, "缺口分两类：", space_after=4)
    add_bullets(doc, [
        "约 110 行：贵司处理后表有，我们的「省二院松北 6 月原始账单」里没有"
        "（常见：宫腔镜包拆开的零件、钉盒类、外来器械 vendor 补录等）",
        "约 117 行：原始文件里其实有（如「盆1碗2盘2/W9050」「上肢钉盒（三）」），"
        "收到完整原始后我们会再排查系统匹配问题",
    ], size=10)
    doc.add_paragraph()

    add_hospital_block(
        doc,
        "4. 黑龙江省第二医院（松北院区）",
        "6 月账单差约 8,743 元。",
        "请提供：",
        [
            "处理后表里有、原始 6 月账单里没有的那部分 —— 对应的原始来源文件",
            "（part3、外来器械、钉盒 vendor 等，贵司内部叫什么名都可以，给文件即可）",
            "像「宫腔镜包」这类整包 → 多行零件的拆包对照（格式同国药，见第一节示例表）",
        ],
    )

    add_para(doc, "四、常见问题", size=14, bold=True, color=COLOR_H2, space_after=6)
    add_para(doc, "问：之前不是已经给过 6 月账单了吗？", bold=True, space_after=4)
    add_para(
        doc,
        "答：之前给的是整份 6 月总账单（原始 + 处理后），材料是齐的。"
        "这次要的是账单里的「拆包明细、分厂家明细、第三部分明细」——"
        "有了这些分录，系统才能和贵司确认版逐行对齐。",
        space_after=10,
    )
    add_para(doc, "问：结款函还需要补吗？", bold=True, space_after=4)
    add_para(
        doc,
        "答：上述 4 家医院的结款函，请等 6 月账单明细对齐后再核对；"
        "目前不必单独再发结款函（除非贵司确认有更新版本）。",
        space_after=14,
    )

    add_para(doc, "附：可转发的一段话（微信/邮件）", size=14, bold=True, color=COLOR_H2, space_after=6)
    quote = (
        "各位老师好，\n\n"
        "特色账单系统 6 月核对进展：37 家医院中，33 家 6 月账单已通过。\n\n"
        "还有 4 家因缺「拆分明细」暂时对不齐，需要补充：\n\n"
        "① 国药总院主院区、第二院区 — 「整包拆零件」对照表（如产包拆成哪些零件、各几件）；\n"
        "② 市二院 — 6 月 7 个供应商/厂家的外来器械补录明细（上海尔欢、大博、尔欢等）；\n"
        "③ 省二松北 — 第三部分/外来器械/钉盒补录 + 拆包对照（如宫腔镜包）。\n\n"
        "说明：6 月总账单我们已有，不必再发；这次要的是「分录级」明细。"
        "Excel 优先，收到一批我们核对一批。如有疑问可约 15 分钟电话，"
        "对着贵司 6 月确认版里多出来的几行具体指给您看。\n\n"
        "详细说明见附件 Word 清单。感谢配合！"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(quote)
    set_run(run, size=10, italic=True)

    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    build()
